# resume_processor.py

import logging
import docx
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from user_interface import UserInterface
from requests.exceptions import RequestException
from rich.progress import Progress, SpinnerColumn, BarColumn, TextColumn, TimeElapsedColumn, TimeRemainingColumn
from rich.console import Console
from rich.panel import Panel
from rich.syntax import Syntax
from difflib import unified_diff

console = Console()

class ResumeProcessor:
    def __init__(self, llm_client):
        self.llm = llm_client

    def load_template(self, file_path):
        try:
            document = docx.Document(file_path)
            return document
        except Exception as e:
            logging.error(f"Failed to load resume template: {str(e)}")
            raise RuntimeError(f"Failed to load resume template: {str(e)}")

    def process_document(self, document, requirements):
        UserInterface.progress("Optimizing resume sections...")
        UserInterface.info("Incorporating key requirements into resume sections...")

        total_sections = self._count_sections(document)
        current_section = 0

        # Define a consistent system message for all sections
        system_message = (
            "You are an expert resume optimizer. Ensure that all responses are accurate, factual, "
            "and adhere strictly to the instructions provided. Do not include any information that "
            "is not present in the input. Your goal is to enhance the resume sections by incorporating "
            "relevant keywords from the key requirements, without increasing the length of the content. "
            "Do not add any emojis to the text. If the original text contains emojis, keep them, but do not introduce new ones."
        )

        with Progress(
            "[progress.description]{task.description}",
            SpinnerColumn(),
            BarColumn(),
            "[progress.percentage]{task.percentage:>3.0f}%",
            "•",
            TimeElapsedColumn(),
            "•",
            TimeRemainingColumn(),
            console=console,
            transient=True
        ) as progress:
            task = progress.add_task("Processing sections...", total=total_sections)

            # Process paragraphs
            for paragraph in document.paragraphs:
                original_text = paragraph.text.strip()
                if original_text:
                    current_section += 1
                    UserInterface.section_status(current_section, total_sections, "paragraph")
                    try:
                        adjusted_text = self.optimize_section(original_text, requirements, system_message)

                        # Print input and output texts with emojis
                        UserInterface.print_input_output(original_text, adjusted_text)

                        # Generate and display diff
                        UserInterface.print_diff(original_text, adjusted_text)

                        # Preserve paragraph style
                        original_style = paragraph.style
                        original_alignment = paragraph.alignment
                        original_format = paragraph.paragraph_format

                        # Preserve runs and their styles, including hyperlinks
                        self._update_paragraph_text(paragraph, adjusted_text)

                        # Apply the original paragraph styles
                        paragraph.style = original_style
                        paragraph.alignment = original_alignment
                        self._copy_paragraph_format(paragraph.paragraph_format, original_format)

                        UserInterface.success(f"Section {current_section}/{total_sections} optimized successfully")
                    except Exception as e:
                        logging.error(f"Error processing paragraph {current_section}: {str(e)}")
                        UserInterface.error(f"Error processing paragraph {current_section}: {str(e)}")
                    progress.update(task, advance=1)

            # Process tables
            for table in document.tables:
                current_section = self._process_table(table, requirements, current_section, total_sections, system_message, progress, task)

    def _process_table(self, table, requirements, current_section, total_sections, system_message, progress, task):
        for row in table.rows:
            for cell in row.cells:
                # Process paragraphs in the cell
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text.strip()
                    if original_text:
                        current_section += 1
                        UserInterface.section_status(current_section, total_sections, "table cell")
                        try:
                            adjusted_text = self.optimize_section(original_text, requirements, system_message)

                            # Print input and output texts with emojis
                            UserInterface.print_input_output(original_text, adjusted_text)

                            # Generate and display diff
                            UserInterface.print_diff(original_text, adjusted_text)

                            # Preserve paragraph style
                            original_style = paragraph.style
                            original_alignment = paragraph.alignment
                            original_format = paragraph.paragraph_format

                            # Preserve runs and their styles, including hyperlinks
                            self._update_paragraph_text(paragraph, adjusted_text)

                            # Apply the original paragraph styles
                            paragraph.style = original_style
                            paragraph.alignment = original_alignment
                            self._copy_paragraph_format(paragraph.paragraph_format, original_format)

                            UserInterface.success(f"Section {current_section}/{total_sections} optimized successfully")
                        except Exception as e:
                            logging.error(f"Error processing table cell {current_section}: {str(e)}")
                            UserInterface.error(f"Error processing table cell {current_section}: {str(e)}")
                        progress.update(task, advance=1)
                # Process nested tables
                for nested_table in cell.tables:
                    current_section = self._process_table(nested_table, requirements, current_section, total_sections, system_message, progress, task)
        return current_section

    def optimize_section(self, content, requirements, system_message):
        # Prepare the prompt
        prompt = (
            f"Resume Section:\n{content}\n\n"
            f"Key Requirements:\n- " + '\n- '.join(requirements) + "\n\n"
            "Instructions:\n"
            "1. Review the section and enhance it by incorporating relevant keywords and phrases from the key requirements if appropriate.\n"
            "2. Do not remove or summarize any existing content; only make necessary adjustments to better match the job requirements.\n"
            "3. Ensure that all original details are preserved exactly as they are.\n"
            "4. If the section is already optimized or no changes are needed, leave it unchanged.\n"
            "5. Keep the output length similar to the input length; do not exceed the original length.\n"
            "6. Do not include any boolean values; all content should be text.\n"
            "7. Preserve special characters and emojis in the text.\n"
            "8. Do not add any emojis to the text. If the original text contains emojis, keep them, but do not introduce new ones.\n"
            "9. Return the optimized section as plain text.\n"
            "Do not include any additional text, explanations, or code block markers."
        )

        try:
            # Estimate max_tokens based on input length
            input_words = len(content.split())
            max_tokens = int(input_words * 1.5)  # Allow up to 50% more tokens

            response = self.llm.generate(prompt, system_message, max_tokens=max_tokens)
            logging.debug(f"LLM Response for section optimization:\n{response}")

            adjusted_content = self._clean_response(response)

            # Ensure that the output length is not significantly longer than the input
            input_length = len(content)
            output_length = len(adjusted_content)
            if output_length > input_length * 1.2:  # Allow up to 20% increase
                logging.warning(f"Adjusted content is significantly longer than input. Truncating to original length.")
                adjusted_content = adjusted_content[:input_length]

            return adjusted_content
        except Exception as e:
            logging.error(f"Exception in optimize_section: {str(e)}")
            raise RuntimeError(f"Failed to optimize section: {str(e)}")

    def _update_paragraph_text(self, paragraph, new_text):
        # This method replaces the paragraph text while preserving hyperlinks and run styles
        # For simplicity, we will attempt to split the new text according to the original runs

        # Get original runs and their properties
        original_runs = []
        for run in paragraph.runs:
            original_runs.append({
                'text': run.text,
                'style': run.style,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'font_color': run.font.color.rgb,
                'hyperlink': self._get_hyperlink(run)
            })

        # Clear the paragraph
        self._clear_paragraph(paragraph)

        # Reconstruct runs with new text
        # For simplicity, we apply the styles of the original runs sequentially to the new text
        # This may not perfectly match the original formatting if the text has changed significantly

        # Combine original run texts to get total length
        original_full_text = ''.join(run['text'] for run in original_runs)
        if not original_full_text:
            return

        # Map positions from original text to new text
        # This is a complex task and may not be perfect
        # For now, we will distribute the new text among the runs proportionally

        new_text_length = len(new_text)
        original_text_length = len(original_full_text)

        if original_text_length == 0:
            return

        # Build new runs
        pos = 0
        for run_info in original_runs:
            run_text = run_info['text']
            run_length = len(run_text)
            proportion = run_length / original_text_length
            new_run_length = int(proportion * new_text_length)
            new_run_text = new_text[pos:pos+new_run_length]
            pos += new_run_length

            if not new_run_text:
                continue

            new_run = paragraph.add_run(new_run_text)
            # Apply styles
            new_run.style = run_info['style']
            new_run.bold = run_info['bold']
            new_run.italic = run_info['italic']
            new_run.underline = run_info['underline']
            new_run.font.name = run_info['font_name']
            new_run.font.size = run_info['font_size']
            new_run.font.color.rgb = run_info['font_color']

            # Handle hyperlinks
            if run_info['hyperlink']:
                self._add_hyperlink(new_run, run_info['hyperlink'])

        # If any remaining text, add it
        if pos < new_text_length:
            remaining_text = new_text[pos:]
            new_run = paragraph.add_run(remaining_text)
            # Apply the style of the last run
            if original_runs:
                last_run_info = original_runs[-1]
                new_run.style = last_run_info['style']
                new_run.bold = last_run_info['bold']
                new_run.italic = last_run_info['italic']
                new_run.underline = last_run_info['underline']
                new_run.font.name = last_run_info['font_name']
                new_run.font.size = last_run_info['font_size']
                new_run.font.color.rgb = last_run_info['font_color']

    def _get_hyperlink(self, run):
        # Check if the run is part of a hyperlink
        # This requires accessing the underlying XML
        try:
            r = run._r
            hlinkClick = r.xpath('./w:rPr/w:hlinkClick')
            if hlinkClick:
                rId = hlinkClick[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                return rId
        except Exception as e:
            pass
        return None

    def _add_hyperlink(self, run, rId):
        # Re-associate the hyperlink relationship to the new run
        hyperlink_rel = run.part.rels[rId]
        hlinkClick = docx.oxml.shared.OxmlElement('w:hlinkClick')
        hlinkClick.set(docx.oxml.shared.qn('r:id'), rId)
        run._r.get_or_add_rPr().append(hlinkClick)

    def _clear_paragraph(self, paragraph):
        p_element = paragraph._element
        p_element.clear_content()

    def _copy_paragraph_format(self, target_format, source_format):
        # Copy paragraph formatting properties individually
        target_format.left_indent = source_format.left_indent
        target_format.right_indent = source_format.right_indent
        target_format.space_before = source_format.space_before
        target_format.space_after = source_format.space_after
        target_format.line_spacing = source_format.line_spacing
        target_format.line_spacing_rule = source_format.line_spacing_rule
        target_format.alignment = source_format.alignment
        target_format.first_line_indent = source_format.first_line_indent
        target_format.keep_together = source_format.keep_together
        target_format.keep_with_next = source_format.keep_with_next
        target_format.page_break_before = source_format.page_break_before
        target_format.widow_control = source_format.widow_control

    def _clean_response(self, response):
        response = response.strip()
        if response.startswith('```') and response.endswith('```'):
            response = response[3:-3].strip()
        if response.startswith('text'):
            response = response[4:].strip()
        return response

    def save_document(self, document, output_path):
        try:
            document.save(output_path)
            UserInterface.success(f"Resume saved as: {output_path}")
        except Exception as e:
            logging.error(f"Failed to save document: {str(e)}")
            raise RuntimeError(f"Failed to save document: {str(e)}")

    def _count_sections(self, document):
        count = 0
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                count += 1
        for table in document.tables:
            count += self._count_sections_in_table(table)
        return count

    def _count_sections_in_table(self, table):
        count = 0
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        count += 1
                for nested_table in cell.tables:
                    count += self._count_sections_in_table(nested_table)
        return count
